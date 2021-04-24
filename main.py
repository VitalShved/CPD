# Control of performance discipline

from docx import Document
from datetime import date, timedelta, datetime
import smtplib
import os
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart



contacts = {
    'Хоменок Ю.В.':'tch10_meh@gomel.rw',
    'Королев С.Н.':'tch10_electr@gomel.rw',
    'Зятиков А.А.':'tch10_snab@gomel.rw',
    'ТЧГ':'tch10_gi@gomel.rw',
    'Швед В.А.':'tch10_tcht@gomel.rw',
    'Ткачев А.С.':'tch10_hoz@gomel.rw',
    'Говор П.В.':'tch10_rmu@gomel.rw',
    'Зезюлин П.В.':'tch10_to@gomel.rw',
    'Говязо Е.А.':'tch10_klad@gomel.rw',
    'Васильцов Д.Г.':'tch10_def@gomel.rw',
    'ТЧЗ-1':'tch10_z1@gomel.rw',
    'Рагина С.М.':'tch10_dom@gomel.rw',
    'Кучеров М.Н.':'tch10_ot@gomel.rw',
    'Дорошенко П.М.':'tch10_nk@gomel.rw',
    'Секретарь':'tch10_tch@gomel.rw'
}


class Parser:

    def __init__(self, filename):
        self.filename = filename

    def process(self):
        result = {}
        content = []
        with open(self.filename, 'rb') as file:
            doc = Document(file)
            for i in range(1, len(doc.tables)):
                if i == 1:
                    for j in range(1, len(doc.tables[i].rows)):
                        line = []
                        for x in range(6):
                            line.append(doc.tables[i].rows[j].cells[x].text)
                        content.append(line)
                else:
                    for j in range(len(doc.tables[i].rows)):
                        line = []
                        for x in range(6):
                            line.append(doc.tables[i].rows[j].cells[x].text)
                        content.append(line)
            for y in range(len(content)):
                result[y] = content[y]
        return result


class Sender:

    def __init__(self, parser, smtp_host, smpt_port, smpt_password):
        self.parser = parser
        self.smtp_host = smtp_host
        self.smtp_port = smpt_port
        self.smtp_password = smpt_password

    def send(self):
        sends = set()
        processed = self.parser.process()
        for value in processed.values():
            sends.add(value[3])
            sends.add(value[4])
        user = 'Skandinav_by@mail.ru'
        recipients = list(sends)
        text = 'For execution'
        filepath = self.parser.filename
        basename = os.path.basename(filepath)
        filesize = os.path.getsize(filepath)
        msg = MIMEMultipart()
        msg['Subject'] = 'Events'
        msg['From'] = 'Skandinav_by@mail.ru'
        msg['To'] = ', '.join(recipients)
        part_text = MIMEText(text, 'plain')
        part_file = MIMEBase('application', 'octet-stream; name="{}"'.format(basename))
        part_file.set_payload(open(filepath, "rb").read())
        part_file.add_header('Content-Description', basename)
        part_file.add_header('Content-Disposition', 'attachment; filename="{}"; size={}'.format(basename, filesize))
        encoders.encode_base64(part_file)
        msg.attach(part_text)
        msg.attach(part_file)
        mail = smtplib.SMTP_SSL(self.smtp_host, self.smtp_port)
        mail.login(user, self.smtp_password)
        mail.sendmail(user, recipients, msg.as_string())
        mail.quit()

    def repeat(self):
        '''Поиск по дате исполнения и повторная рассылка приорететных задач'''
        processed = self.parser.process()
        for value in processed.values():
            day = datetime.strptime(value[2], '%d.%m.%Y').date()
            date_today = date.today()
            new_date = date_today
            new_date1 = date_today + timedelta(days=3)
            new_date2 = date_today + timedelta(days=1)
            if day == new_date1:
                for recipients in contacts.keys():
                    if value[3] == recipients:
                        user = 'Skandinav_by@mail.ru'
                        text = f'Повторно напоминаем о выполнении пункта мероприятий {value}'
                        msg = MIMEMultipart()
                        msg['Subject'] = 'Events'
                        msg['From'] = 'Skandinav_by@mail.ru'
                        msg['To'] = contacts.get(recipients)
                        part_text = MIMEText(text, 'plain')
                        msg.attach(part_text)
                        mail = smtplib.SMTP_SSL(self.smtp_host, self.smtp_port)
                        mail.login(user, self.smtp_password)
                        mail.sendmail(user, recipients, msg.as_string())
                        mail.quit()
                    elif day == new_date2:
                        if value[3] == recipients:
                                user = 'Skandinav_by@mail.ru'
                                text = f'Повторно напоминаем о выполнении пункта мероприятий {value}'
                                msg = MIMEMultipart()
                                msg['Subject'] = 'Events'
                                msg['From'] = 'Skandinav_by@mail.ru'
                                msg['To'] = contacts.get(recipients)
                                part_text = MIMEText(text, 'plain')
                                msg.attach(part_text)
                                mail = smtplib.SMTP_SSL(self.smtp_host, self.smtp_port)
                                mail.login(user, self.smtp_password)
                                mail.sendmail(user, recipients, msg.as_string())
                                mail.quit()
                        if value[4] == recipients:
                                user = 'Skandinav_by@mail.ru'
                                text = f'Повторно напоминаем о выполнении пункта мероприятий {value}'
                                msg = MIMEMultipart()
                                msg['Subject'] = 'Events'
                                msg['From'] = 'Skandinav_by@mail.ru'
                                msg['To'] = contacts.get(recipients)
                                part_text = MIMEText(text, 'plain')
                                msg.attach(part_text)
                                mail = smtplib.SMTP_SSL(self.smtp_host, self.smtp_port)
                                mail.login(user, self.smtp_password)
                                mail.sendmail(user, recipients, msg.as_string())
                                mail.quit()
                    elif day < new_date:
                        user = 'Skandinav_by@mail.ru'
                        text = f'Не были выполнены данные мероприятия {value}'
                        msg = MIMEMultipart()
                        msg['Subject'] = 'Events'
                        msg['From'] = 'Skandinav_by@mail.ru'
                        msg['To'] = contacts.get("Секретарь")
                        part_text = MIMEText(text, 'plain')
                        msg.attach(part_text)
                        mail = smtplib.SMTP_SSL(self.smtp_host, self.smtp_port)
                        mail.login(user, self.smtp_password)
                        mail.sendmail(user, recipients, msg.as_string())
                        mail.quit()




parser = Parser('events.docx')
# print(parser.process())
sender = Sender(parser, 'smtp.mail.ru', 25, 'password???')

sender.send()
sender.repeat()

